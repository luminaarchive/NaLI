"""
NaLI Evidence Report — Journal-Grade DOCX Generator
Standard: Animal Conservation / Wiley academic layout
"""

import io
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Color palette ─────────────────────────────────────────────────────────
def rgb(hex_str: str) -> RGBColor:
    h = hex_str.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))

COL_GREEN_DARK  = rgb("1B5E20")
COL_GREEN_MED   = rgb("2E7D32")
COL_TEAL_MED    = rgb("00897B")
COL_AMBER       = rgb("F57C00")
COL_ROSE        = rgb("C62828")
COL_WHITE       = rgb("FFFFFF")
COL_NEAR_BLACK  = rgb("121212")
COL_GRAY_MID    = rgb("757575")
COL_GRAY_DARK   = rgb("424242")
COL_SCORE_BG    = rgb("E0F2F1")
COL_INFO_BG     = rgb("FAFAFA")
COL_ROW_ALT     = rgb("F9FBE7")
COL_Q_FILL      = rgb("E0F2F1")
COL_GREEN_PALE  = rgb("E8F5E9")
COL_AMBER_PALE  = rgb("FFF8E1")
COL_TEAL_BG     = rgb("F0FAF7")
COL_LG_BG       = rgb("F5F5F5")


def score_color(score: int) -> RGBColor:
    if score >= 80:
        return COL_TEAL_MED
    if score >= 60:
        return COL_AMBER
    if score >= 40:
        return rgb("F97316")
    return COL_ROSE


# ── Low-level XML helpers ─────────────────────────────────────────────────
def set_cell_bg(cell, hex_color: str):
    """Fill a table cell background color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tcPr.append(shd)


def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    """Set individual cell borders. Each value is (size_8ths_pt, color_hex)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")

    def make_border(side, size, color):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color.lstrip("#"))
        return el

    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if val:
            tcBorders.append(make_border(side, val[0], val[1]))
    tcPr.append(tcBorders)


def no_borders(table):
    """Remove all borders from a table."""
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        tblBorders.append(el)
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(tblBorders)


def thin_borders(table, color_hex: str = "BDBDBD"):
    """Apply thin grid borders to a table."""
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color_hex.lstrip("#"))
        tblBorders.append(el)
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(tblBorders)


def set_col_widths(table, widths_cm: list):
    """Set column widths in cm."""
    tbl = table._tbl
    tblGrid = tbl.find(qn("w:tblGrid"))
    if tblGrid is None:
        tblGrid = OxmlElement("w:tblGrid")
        tbl.insert(0, tblGrid)
    else:
        for child in list(tblGrid):
            tblGrid.remove(child)
    for w in widths_cm:
        gridCol = OxmlElement("w:gridCol")
        gridCol.set(qn("w:w"), str(int(w * 567)))  # 1cm = 567 twips
        tblGrid.append(gridCol)


def cell_padding(cell, top_pt=3, bottom_pt=3, left_pt=3, right_pt=3):
    """Set cell padding."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top", top_pt), ("bottom", bottom_pt), ("left", left_pt), ("right", right_pt)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(int(val * 20)))  # pt to twips
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)


def set_vertical_align(cell, align="top"):
    """Set cell vertical alignment."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement("w:vAlign")
    vAlign.set(qn("w:val"), align)
    tcPr.append(vAlign)


def add_run_colored(para, text: str, bold=False, italic=False,
                    color=None, size_pt=None, font_name=None):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if size_pt:
        run.font.size = Pt(size_pt)
    if font_name:
        run.font.name = font_name
    return run


def add_bottom_border(para, color_hex: str = "2E7D32"):
    """Add a bottom border to a paragraph (section heading style)."""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex.lstrip("#"))
    pBdr.append(bottom)
    pPr.append(pBdr)


# ── Header / Footer helpers ───────────────────────────────────────────────
def add_header(doc, data: dict):
    """Add header to all body sections (skip cover)."""
    score = data.get("score", 0)
    level = data.get("level", "")
    date = data.get("date", "")
    header = doc.sections[0].header
    para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    para.clear()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(f"NaLI Evidence Report  |  {date}  |  Score: {score}% — {level}")
    run.bold = True
    run.font.size = Pt(8.5)
    run.font.color.rgb = COL_GREEN_DARK


def add_footer(doc, data: dict):
    """Add footer to all body sections."""
    date = data.get("date", "")
    footer = doc.sections[0].footer
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.clear()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(
        "Draft NaLI · Pemeriksaan akhir dan tanggung jawab ilmiah tetap pada pengguna · naliai.vercel.app"
    )
    run.italic = True
    run.font.size = Pt(7.5)
    run.font.color.rgb = COL_GRAY_MID

    # Add page number field to right
    tab = para.add_run("\t")
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.text = " PAGE "
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run2 = para.add_run()
    run2._r.append(fldChar1)
    run2._r.append(instrText)
    run2._r.append(fldChar2)
    run2.font.size = Pt(7.5)
    run2.font.color.rgb = COL_GRAY_MID


# ── Section heading ───────────────────────────────────────────────────────
def add_section_heading(doc, text: str):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run(text.upper())
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.color.rgb = COL_GREEN_DARK
    add_bottom_border(para)
    return para


# ── Body paragraph with inline labels ────────────────────────────────────
def add_body_para(doc, text: str):
    """Add a body paragraph with color-coded inline labels."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.space_after = Pt(5)

    parts = []
    buf = ""
    i = 0
    while i < len(text):
        if text[i:].startswith("[Terkonfirmasi]"):
            if buf:
                parts.append(("normal", buf))
                buf = ""
            parts.append(("ok", "[Terkonfirmasi]"))
            i += len("[Terkonfirmasi]")
        elif text[i:].startswith("[Inferensi AI]"):
            if buf:
                parts.append(("normal", buf))
                buf = ""
            parts.append(("infer", "[Inferensi AI]"))
            i += len("[Inferensi AI]")
        elif text[i:].startswith("[Bukti kurang]"):
            if buf:
                parts.append(("normal", buf))
                buf = ""
            parts.append(("weak", "[Bukti kurang]"))
            i += len("[Bukti kurang]")
        else:
            buf += text[i]
            i += 1
    if buf:
        parts.append(("normal", buf))

    for kind, segment in parts:
        run = para.add_run(segment)
        run.font.name = "Times New Roman"
        run.font.size = Pt(10.5)
        if kind == "ok":
            run.bold = True
            run.font.color.rgb = COL_TEAL_MED
        elif kind == "infer":
            run.bold = True
            run.font.color.rgb = COL_AMBER
        elif kind == "weak":
            run.bold = True
            run.font.color.rgb = COL_ROSE
        else:
            run.font.color.rgb = COL_NEAR_BLACK

    return para


# ── Cover page ───────────────────────────────────────────────────────────
def build_cover(doc, data: dict):
    """Build the cover page with green header, score box, and article info."""
    score = data.get("score", 0)
    level = data.get("level", "")
    title = data.get("title", "NaLI Report")
    tipe = data.get("tipe", "Laporan")
    kualitas = data.get("kualitas", "Sedang")
    risiko = data.get("risiko", "Sedang")
    date = data.get("date", "")
    g = data.get("g", 0.0)
    v = data.get("v", 0.0)
    h = data.get("h", 0.0)
    i_val = data.get("i", 0.0)
    li = data.get("li", 1.0)
    decay = data.get("decay", 1.0)
    keywords = data.get("keywords", "")
    sc = score_color(score)

    # Green header table (full width, 1 row)
    hdr_tbl = doc.add_table(rows=1, cols=1)
    no_borders(hdr_tbl)
    cell = hdr_tbl.rows[0].cells[0]
    set_cell_bg(cell, "1B5E20")
    cell_padding(cell, 10, 10, 12, 12)

    p1 = cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run("NaLI EVIDENCE REPORT")
    r1.bold = True
    r1.font.size = Pt(20)
    r1.font.color.rgb = COL_WHITE

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Nature Life Intelligence and Human Assistance")
    r2.italic = True
    r2.font.size = Pt(11)
    r2.font.color.rgb = rgb("E8F5E9")

    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("naliai.vercel.app  |  Evidence Intelligence OS v2.0")
    r3.font.size = Pt(9)
    r3.font.color.rgb = rgb("A5D6A7")

    doc.add_paragraph()  # small spacer

    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title)
    title_run.bold = True
    title_run.font.name = "Times New Roman"
    title_run.font.size = Pt(16)
    title_run.font.color.rgb = COL_NEAR_BLACK

    doc.add_paragraph()

    # 2-column table: score+pillars | article info
    info_tbl = doc.add_table(rows=1, cols=2)
    no_borders(info_tbl)
    set_col_widths(info_tbl, [7.0, 10.5])

    left_cell = info_tbl.rows[0].cells[0]
    right_cell = info_tbl.rows[0].cells[1]
    set_cell_bg(left_cell, "E0F2F1")
    set_cell_bg(right_cell, "FAFAFA")
    cell_padding(left_cell, 8, 8, 8, 8)
    cell_padding(right_cell, 8, 8, 8, 8)
    set_vertical_align(left_cell, "top")
    set_vertical_align(right_cell, "top")

    # Left: score
    sp = left_cell.paragraphs[0]
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sp.add_run(f"{score}%")
    sr.bold = True
    sr.font.size = Pt(36)
    sr.font.color.rgb = sc

    sp2 = left_cell.add_paragraph()
    sp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr2 = sp2.add_run(level)
    sr2.bold = True
    sr2.font.size = Pt(9)
    sr2.font.color.rgb = sc

    sp3 = left_cell.add_paragraph()
    sp3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr3 = sp3.add_run("PALANTIR CONFIDENCE SCORE")
    sr3.font.size = Pt(7)
    sr3.font.color.rgb = COL_GRAY_MID

    # Pillar rows
    for label, val in [
        (f"Genetik x0.60", f"{g:.2f}"),
        (f"Visual x0.20", f"{v:.2f}"),
        (f"Habitat x0.15", f"{h:.2f}"),
        (f"Integritas x0.10", f"{i_val:.2f}"),
        (f"Li Multiplier", f"{li:.2f}x"),
        (f"Temporal P(t)", f"{decay:.2f}"),
    ]:
        pp = left_cell.add_paragraph()
        pp.paragraph_format.space_after = Pt(1)
        rl = pp.add_run(f"{label}: ")
        rl.font.size = Pt(8)
        rl.font.color.rgb = COL_GRAY_DARK
        rv = pp.add_run(val)
        rv.bold = True
        rv.font.size = Pt(8)
        rv.font.color.rgb = COL_NEAR_BLACK

    # Right: article info
    rp0 = right_cell.paragraphs[0]
    rp0.paragraph_format.space_after = Pt(3)
    rr0 = rp0.add_run("INFORMASI LAPORAN")
    rr0.bold = True
    rr0.font.size = Pt(9)
    rr0.font.color.rgb = COL_GRAY_DARK
    add_bottom_border(rp0, "BDBDBD")

    rp1 = right_cell.add_paragraph()
    rp1.paragraph_format.space_after = Pt(3)
    rr1 = rp1.add_run(tipe)
    rr1.bold = True
    rr1.font.size = Pt(8.5)
    rr1.font.color.rgb = COL_NEAR_BLACK

    rp2 = right_cell.add_paragraph()
    rp2.paragraph_format.space_after = Pt(2)
    add_run_colored(rp2, "Kualitas Bukti: ", size_pt=8, color=COL_GRAY_MID)
    add_run_colored(rp2, kualitas, bold=True, size_pt=10, color=sc)

    rp3 = right_cell.add_paragraph()
    rp3.paragraph_format.space_after = Pt(2)
    add_run_colored(rp3, "Risiko Klaim: ", size_pt=8, color=COL_GRAY_MID)
    add_run_colored(rp3, risiko, bold=True, size_pt=10, color=COL_AMBER)

    rp4 = right_cell.add_paragraph()
    rp4.paragraph_format.space_after = Pt(2)
    add_run_colored(rp4, f"Tanggal: {date}", size_pt=8, color=COL_GRAY_DARK)

    rp5 = right_cell.add_paragraph()
    add_run_colored(rp5, "Sistem: NaLI OS v2.0", size_pt=8, color=COL_GRAY_DARK)

    # Keywords row
    if keywords:
        doc.add_paragraph()
        kw_para = doc.add_paragraph()
        kw_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        kw_run = kw_para.add_run(f"Kata Kunci: {keywords}")
        kw_run.italic = True
        kw_run.font.size = Pt(9)
        kw_run.font.color.rgb = COL_GRAY_MID


# ── Abstract section ──────────────────────────────────────────────────────
def build_abstract(doc, data: dict):
    abs_id = data.get("abstract_id", "")
    abs_en = data.get("abstract_en", "")

    if not abs_id and not abs_en:
        return

    abs_tbl = doc.add_table(rows=1, cols=1)
    no_borders(abs_tbl)
    cell = abs_tbl.rows[0].cells[0]
    set_cell_bg(cell, "FAFAFA")
    set_cell_border(cell, top=(4, "BDBDBD"), bottom=(4, "BDBDBD"),
                    left=(4, "BDBDBD"), right=(4, "BDBDBD"))
    cell_padding(cell, 8, 8, 10, 10)

    p0 = cell.paragraphs[0]
    r0 = p0.add_run("ABSTRAK")
    r0.bold = True
    r0.font.size = Pt(9)
    r0.font.color.rgb = COL_NEAR_BLACK

    if abs_id:
        p1 = cell.add_paragraph()
        p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p1.paragraph_format.space_after = Pt(4)
        r1 = p1.add_run(abs_id)
        r1.italic = True
        r1.font.name = "Times New Roman"
        r1.font.size = Pt(9)

    if abs_en:
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_before = Pt(4)
        r2a = p2.add_run("Abstract")
        r2a.bold = True
        r2a.italic = True
        r2a.font.size = Pt(9)

        p3 = cell.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r3 = p3.add_run(abs_en)
        r3.italic = True
        r3.font.name = "Times New Roman"
        r3.font.size = Pt(9)


# ── Evidence table ────────────────────────────────────────────────────────
def build_evidence_table(doc, ev_table: list):
    if not ev_table or len(ev_table) < 2:
        return

    col_widths = [6.5, 2.5, 2.8, 5.2]
    n_rows = len(ev_table)
    tbl = doc.add_table(rows=n_rows, cols=4)
    thin_borders(tbl, "BDBDBD")
    set_col_widths(tbl, col_widths)

    header = ev_table[0] if ev_table else ["Klaim", "Pilar Bukti", "Status", "Catatan"]
    header_row = tbl.rows[0]
    for j, h_text in enumerate(header[:4]):
        cell = header_row.cells[j]
        set_cell_bg(cell, "1B5E20")
        cell_padding(cell, 3, 3, 3, 3)
        set_vertical_align(cell, "center")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(str(h_text))
        r.bold = True
        r.font.size = Pt(8)
        r.font.color.rgb = COL_WHITE

    for i, row_data in enumerate(ev_table[1:]):
        cells = list(row_data) + [""] * max(0, 4 - len(row_data))
        bg_hex = "F9FBE7" if i % 2 == 0 else "FFFFFF"
        tbl_row = tbl.rows[i + 1]
        for j in range(4):
            cell = tbl_row.cells[j]
            set_cell_bg(cell, bg_hex)
            cell_padding(cell, 3, 3, 3, 3)
            set_vertical_align(cell, "top")
            p = cell.paragraphs[0]

            if j == 2:
                status = str(cells[j])
                s = status.lower()
                r = p.add_run(status)
                r.bold = True
                r.font.size = Pt(8)
                if "terkonfirmasi" in s:
                    r.font.color.rgb = COL_TEAL_MED
                elif "inferensi" in s:
                    r.font.color.rgb = COL_AMBER
                else:
                    r.font.color.rgb = COL_ROSE
            else:
                r = p.add_run(str(cells[j]))
                r.font.size = Pt(8.5)
                r.font.color.rgb = COL_NEAR_BLACK


# ── Missing evidence ──────────────────────────────────────────────────────
def build_missing_evidence(doc, missing: list):
    for item in missing:
        num = str(item[0]) if len(item) > 0 else "?"
        name = str(item[1]) if len(item) > 1 else ""
        impact = str(item[2]) if len(item) > 2 else ""

        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(3)
        rn = para.add_run(f"{num}. ")
        rn.bold = True
        rn.font.size = Pt(9)
        rn.font.color.rgb = COL_TEAL_MED
        rname = para.add_run(name)
        rname.bold = True
        rname.font.size = Pt(9)
        rname.font.color.rgb = COL_NEAR_BLACK
        if impact:
            para.add_run("  ")
            rimp = para.add_run(f"— {impact}")
            rimp.italic = True
            rimp.font.size = Pt(8.5)
            rimp.font.color.rgb = COL_GRAY_MID


# ── Follow-up questions ───────────────────────────────────────────────────
def build_questions(doc, questions: list):
    labels = ["Q1", "Q2", "Q3", "Q4", "Q5"]
    for i, q in enumerate(questions):
        label = labels[i] if i < len(labels) else f"Q{i+1}"
        tbl = doc.add_table(rows=1, cols=2)
        no_borders(tbl)
        set_col_widths(tbl, [1.2, 15.8])

        left = tbl.rows[0].cells[0]
        right = tbl.rows[0].cells[1]
        set_cell_bg(left, "E0F2F1")
        set_cell_bg(right, "E0F2F1")
        set_cell_border(left, top=(4, "B2DFDB"), bottom=(4, "B2DFDB"),
                        left=(4, "B2DFDB"), right=(2, "B2DFDB"))
        set_cell_border(right, top=(4, "B2DFDB"), bottom=(4, "B2DFDB"),
                        left=(2, "B2DFDB"), right=(4, "B2DFDB"))
        cell_padding(left, 5, 5, 6, 6)
        cell_padding(right, 5, 5, 6, 6)
        set_vertical_align(left, "center")
        set_vertical_align(right, "top")

        lp = left.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lr = lp.add_run(label)
        lr.bold = True
        lr.font.size = Pt(9)
        lr.font.color.rgb = COL_TEAL_MED

        rp = right.paragraphs[0]
        rp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        rr = rp.add_run(str(q))
        rr.font.size = Pt(10)
        rr.font.color.rgb = COL_NEAR_BLACK

        doc.add_paragraph().paragraph_format.space_after = Pt(2)


# ── Integrity statement ───────────────────────────────────────────────────
def build_integrity(doc, data: dict):
    score = data.get("score", 0)
    level = data.get("level", "")
    date = data.get("date", "")

    it = doc.add_table(rows=1, cols=1)
    no_borders(it)
    cell = it.rows[0].cells[0]
    set_cell_bg(cell, "E8F5E9")
    cell_padding(cell, 8, 8, 10, 10)

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        f"Draft NaLI · Evidence-Grade Intelligence OS v2.0 · "
        f"Palantir Score: {score}% ({level}) · naliai.vercel.app · {date} · "
        f"Pemeriksaan akhir dan tanggung jawab ilmiah tetap pada pengguna."
    )
    r.italic = True
    r.font.size = Pt(8)
    r.font.color.rgb = COL_GREEN_DARK


# ── Main builder ──────────────────────────────────────────────────────────
def build_docx(output, data: dict) -> None:
    """Build DOCX to file path (str) or BytesIO buffer."""
    doc = Document()

    # Page margins
    section = doc.sections[0]
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    add_header(doc, data)
    add_footer(doc, data)

    # ─── COVER PAGE ──────────────────────────────────────────────────
    build_cover(doc, data)

    # Page break after cover
    doc.add_page_break()

    # ─── ABSTRACT ────────────────────────────────────────────────────
    build_abstract(doc, data)
    doc.add_paragraph()

    # ─── BODY SECTIONS ───────────────────────────────────────────────
    for heading, body in data.get("sections", []):
        add_section_heading(doc, str(heading))
        for line in str(body).strip().split("\n"):
            line = line.strip()
            if not line:
                continue
            add_body_para(doc, line)

    # ─── EVIDENCE TABLE ──────────────────────────────────────────────
    ev_table = data.get("ev_table", [])
    if ev_table and len(ev_table) > 1:
        add_section_heading(doc, "TABEL BUKTI DAN KLAIM")
        build_evidence_table(doc, ev_table)
        doc.add_paragraph()

    # ─── MISSING EVIDENCE ────────────────────────────────────────────
    missing = data.get("missing", [])
    if missing:
        add_section_heading(doc, "BUKTI YANG MASIH DIBUTUHKAN")
        build_missing_evidence(doc, missing)
        doc.add_paragraph()

    # ─── FOLLOW-UP QUESTIONS ─────────────────────────────────────────
    questions = data.get("questions", [])
    if questions:
        add_section_heading(doc, "PERTANYAAN LANJUTAN DARI NALI")
        intro = doc.add_paragraph()
        intro_run = intro.add_run(
            "Pertanyaan ini dapat membantu meningkatkan kualitas laporan Anda."
        )
        intro_run.italic = True
        intro_run.font.size = Pt(8.5)
        intro_run.font.color.rgb = COL_GRAY_MID
        intro.paragraph_format.space_after = Pt(4)
        build_questions(doc, questions)

    # ─── REFERENCES ──────────────────────────────────────────────────
    refs = data.get("refs", [])
    if refs:
        add_section_heading(doc, "REFERENSI")
        for ref in refs:
            rp = doc.add_paragraph()
            rp.paragraph_format.space_after = Pt(4)
            rp.paragraph_format.left_indent = Cm(0.5)
            rp.paragraph_format.first_line_indent = Cm(-0.5)
            rr = rp.add_run(str(ref))
            rr.font.name = "Times New Roman"
            rr.font.size = Pt(9.5)
            rr.font.color.rgb = COL_NEAR_BLACK

    # ─── INTEGRITY STATEMENT ─────────────────────────────────────────
    doc.add_paragraph()
    build_integrity(doc, data)

    if hasattr(output, "write"):
        doc.save(output)
    else:
        doc.save(output)


def build_docx_bytes(data: dict) -> bytes:
    """Return DOCX as bytes."""
    buf = io.BytesIO()
    build_docx(buf, data)
    return buf.getvalue()


# ── Self-test & CLI entry point ───────────────────────────────────────────
TEST_DATA = {
    "title": "Observasi Elang Jawa (Nisaetus bartelsi) di Lereng Gunung Semeru: Draft Laporan Berbasis Bukti",
    "subtitle": "Draft Laporan Berbasis Bukti — NaLI Evidence Intelligence OS v2.0",
    "date": "31 Mei 2026",
    "score": 43, "level": "INDIKASI AWAL",
    "kualitas": "Sedang", "risiko": "Sedang",
    "tipe": "Laporan Observasi Satwa",
    "g": 0.00, "v": 0.20, "h": 0.80, "i": 0.45,
    "li": 1.15, "decay": 1.00,
    "abstract_id": "Laporan ini menyajikan draft awal hasil observasi lapangan terhadap satu individu Elang Jawa (Nisaetus bartelsi) di lereng Gunung Semeru, Jawa Timur, pada ketinggian 1.800 mdpl.",
    "abstract_en": "This report presents a preliminary evidence-based draft documenting a single adult male Javan Hawk-Eagle (Nisaetus bartelsi) at 1,800 m a.s.l. on Mount Semeru, East Java, on 12 May 2026.",
    "keywords": "Nisaetus bartelsi, Elang Jawa, Gunung Semeru, evidence-based report, NaLI",
    "sections": [
        ("PENDAHULUAN", "Elang Jawa (Nisaetus bartelsi) merupakan raptor endemik Pulau Jawa yang terdaftar sebagai Endangered (EN) dalam Daftar Merah IUCN."),
        ("METODE PENGAMATAN", "[Bukti kurang] Metode tidak dijelaskan secara eksplisit. Observasi visual langsung selama 2 jam tanpa protokol terstandar."),
        ("HASIL OBSERVASI", "Satu individu dewasa jantan [Terkonfirmasi] teramati pada ketinggian 1.800 mdpl. [Inferensi AI] Perilaku konsisten dengan aktivitas pagi raptor puncak."),
        ("ANALISIS", "Lokasi konsisten dengan known range. [Inferensi AI] Mungkin home range aktif."),
        ("KETERBATASAN DATA", "1. Tidak ada foto/video.\n2. Tidak ada sampel genetik.\n3. Satu kali observasi."),
        ("KESIMPULAN", "[Terkonfirmasi] Keberadaan Nisaetus bartelsi dikonfirmasi di lokasi. Score 43% memerlukan data tambahan."),
        ("REKOMENDASI", "1. Pasang kamera trap.\n2. Kumpulkan eDNA.\n3. Catat GPS eksak."),
    ],
    "ev_table": [
        ["Klaim", "Pilar Bukti", "Status", "Catatan"],
        ["Keberadaan N. bartelsi", "Visual", "Terkonfirmasi", "1 individu, 1 kali"],
        ["Habitat hutan primer 80%", "Deskripsi", "Terkonfirmasi", "Estimasi visual"],
        ["Perilaku berburu", "Perilaku", "Inferensi AI", "Tanpa rekaman"],
        ["Home range aktif", "Ekologi", "Bukti kurang", "Butuh telemetri"],
    ],
    "missing": [
        ("1", "Sampel genetik / eDNA", "+35 poin — Pilar Genetik 0.00 ke 0.85"),
        ("2", "Foto atau video", "+12 poin — Pilar Visual 0.20 ke 0.80"),
        ("3", "GPS koordinat eksak", "Untuk kalibrasi habitat"),
    ],
    "questions": [
        "Apakah tersedia foto atau video dari pengamatan ini?",
        "Di koordinat GPS mana tepatnya observasi dilakukan?",
        "Apakah ada tanda sarang atau individu lain di sekitar lokasi?",
    ],
    "refs": [
        "van Balen, S., Nijboer, J., & Meyburg, B.U. (1999). Distribution and Conservation of the Javan Hawk-eagle. Journal of Raptor Research, 33(2), 105-116.",
        "BirdLife International (2023). Nisaetus bartelsi. IUCN Red List 2023.",
        "NaLI Evidence Intelligence OS (2026). Palantir Scoring Framework v2.0. NatIve, Semarang.",
    ],
}

if __name__ == "__main__":
    import sys
    import json as _json
    import os

    if len(sys.argv) == 3:
        input_file = sys.argv[1]
        output_file = sys.argv[2]
        with open(input_file, "r", encoding="utf-8") as f:
            data = _json.load(f)
        build_docx(output_file, data)
        print(f"OK:{output_file}")
    else:
        OUT = "/tmp/nali_test_output.docx"
        build_docx(OUT, TEST_DATA)

        from docx import Document as DocxReader
        doc = DocxReader(OUT)
        # Collect text from paragraphs AND table cells
        all_texts = [p.text for p in doc.paragraphs]
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    all_texts.append(cell.text)
        full_text = " ".join(all_texts)
        assert "NaLI EVIDENCE REPORT" in full_text, f"Cover title missing. Got: {full_text[:300]}"
        assert any("PENDAHULUAN" in p.text for p in doc.paragraphs), "Body sections missing"

        size = os.path.getsize(OUT)
        assert size > 20000, f"DOCX too small ({size} bytes)"

        print(f"DOCX SELF-TEST PASSED: {len(doc.paragraphs)} paragraphs, {size/1024:.1f}KB")
        print(f"Output: {OUT}")
